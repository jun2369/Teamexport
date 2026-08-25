import base64
import io

from docx import Document
from docx.shared import Inches


def build_docx(messages, chat_topic):
    doc = Document()
    doc.add_heading(chat_topic, level=1)

    for msg in messages:
        header = doc.add_paragraph()
        run = header.add_run(f"{msg['sender']}  —  {msg['created']}")
        run.bold = True

        if msg["text"]:
            doc.add_paragraph(msg["text"])

        for image in msg["images"]:
            img_bytes = base64.b64decode(image["bytes_b64"])
            try:
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(4))
            except Exception:
                doc.add_paragraph("[image could not be embedded]")

        for att in msg.get("attachments", []):
            doc.add_paragraph(f"\U0001F4CE {att['name']}: {att['url']}")

        doc.add_paragraph("")  # spacing between messages

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_txt(messages, chat_topic):
    lines = [chat_topic, "=" * len(chat_topic), ""]

    for msg in messages:
        lines.append(f"{msg['sender']}  —  {msg['created']}")
        if msg["text"]:
            lines.append(msg["text"])
        if msg["images"]:
            lines.append(f"[{len(msg['images'])} image(s) omitted from TXT export]")
        for att in msg.get("attachments", []):
            lines.append(f"[Attachment] {att['name']}: {att['url']}")
        lines.append("")

    buf = io.BytesIO("\n".join(lines).encode("utf-8"))
    buf.seek(0)
    return buf
